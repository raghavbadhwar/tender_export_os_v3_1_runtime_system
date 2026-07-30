from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    "/Users/raghav/Downloads/tender_export_os_v3_1_runtime_system/"
    "outputs/partner_briefs/Tender_Export_OS_Partner_Brief_Grant_Thornton.docx"
)

# Resolved preset: standard_business_brief.
# Named visual overrides used consistently for the partner-pack opening block.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PARTNER_NAVY = "163A5F"
PARTNER_GOLD = "B78328"
BODY = "222222"
MUTED = "5C6670"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9DEE5"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[side]))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_borders(table, color: str = BORDER, size: int = 6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    set_run_font(run, size=9, color=MUTED)


def set_paragraph_shading_and_border(paragraph, *, fill: str, border_color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")

    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)


def add_custom_numbering(document: Document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    existing_abstract = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))
    ]
    abstract_id = (max(existing_abstract) + 1) if existing_abstract else 1
    num_id = (max(existing_num) + 1) if existing_num else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    lvl.append(r_pr)

    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_body(doc: Document, text: str, *, after: float = 6, keep_with_next=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep_with_next
    r = p.add_run(text)
    set_run_font(r, size=11, color=BODY)
    return p


def add_labeled_paragraph(
    doc: Document,
    label: str,
    text: str,
    *,
    after: float = 6,
    indent: float = 0,
):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r1 = p.add_run(label)
    set_run_font(r1, size=11, color=PARTNER_NAVY, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=BODY)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_numbered_step(
    doc: Document, num_id: int, title: str, text: str
):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    r1 = p.add_run(title)
    set_run_font(r1, size=11, color=PARTNER_NAVY, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=BODY)
    return p


def add_bullet(doc: Document, num_id: int, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=PARTNER_NAVY, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=BODY)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=BODY)
    return p


def configure_section_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def clear_paragraph_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def populate_page_furniture(section):
    header = section.header
    hp = header.paragraphs[0]
    clear_paragraph_runs(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT
    )
    left = hp.add_run("TENDER EXPORT OS")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    hp.add_run("\t")
    right = hp.add_run("PARTNER BRIEF  |  JULY 2026")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph_runs(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    label = fp.add_run("Page ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(fp)


def add_partner_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_geometry(section)
    # python-docx omits w:type for the default next-page section. Word accepts
    # that, but LibreOffice can treat the intermediate section inconsistently;
    # encode the transition explicitly for deterministic rendering.
    sect_pr = section._sectPr
    section_type = sect_pr.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        sect_pr.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    populate_page_furniture(section)
    return section


def style_document(doc: Document):
    section = doc.sections[0]
    configure_section_geometry(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Quiet multi-page partner-pack furniture.
    populate_page_furniture(section)


def build():
    doc = Document()
    style_document(doc)
    numbered_id = add_custom_numbering(doc, bullet=False)
    bullet_id = add_custom_numbering(doc, bullet=True)

    # Customer-pack opening block.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("AUTOMATION & OPERATING MODEL BRIEF")
    set_run_font(r, size=10, color=PARTNER_GOLD, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    tr = title.add_run("Tender Export OS")
    set_run_font(tr, size=29, color=PARTNER_NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.keep_with_next = True
    sr = subtitle.add_run(
        "How I turned a fragmented tender process into a governed, "
        "evidence-led operating system"
    )
    set_run_font(sr, size=13.5, color=MUTED)

    add_labeled_paragraph(
        doc,
        "Prepared for: ",
        "Partner and Head of Automation, Grant Thornton",
        after=2,
    )
    add_labeled_paragraph(doc, "Prepared by: ", "Raghav", after=2)
    add_labeled_paragraph(
        doc, "Date: ", date(2026, 7, 30).strftime("%d %B %Y"), after=14
    )

    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(14)
    callout.paragraph_format.left_indent = Inches(0.16)
    callout.paragraph_format.right_indent = Inches(0.10)
    callout.paragraph_format.line_spacing = 1.12
    set_paragraph_shading_and_border(
        callout, fill=LIGHT_BLUE, border_color=BLUE
    )
    cr1 = callout.add_run("In one line: ")
    set_run_font(cr1, size=11, color=PARTNER_NAVY, bold=True)
    cr2 = callout.add_run(
        "I have built a supervised automation system that helps a tender desk "
        "find, understand, qualify and prepare opportunities—while keeping "
        "commercially or legally binding decisions in human hands."
    )
    set_run_font(cr2, size=11, color=BODY)

    add_heading(doc, "Why I built it", 1)
    add_body(
        doc,
        "Tendering is rarely one clean process. An opportunity may begin on a "
        "portal, move into a PDF or BOQ, require supplier discovery, pass through "
        "pricing and compliance checks, and finally depend on somebody remembering "
        "a deadline or approval. In practice, this work is usually spread across "
        "browser tabs, spreadsheets, emails, folders and individual memory.",
    )
    add_body(
        doc,
        "I wanted to create something broader than a tender scraper. The idea was "
        "to build the operating backbone of a tender and export desk: every "
        "opportunity receives a case ID, evidence, a current status, a clear next "
        "step and an audit trail. The system assists the team throughout the "
        "process, but it does not pretend that AI should make every decision.",
        after=10,
    )

    add_heading(doc, "What the system does", 1)
    flow = doc.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow.paragraph_format.space_before = Pt(4)
    flow.paragraph_format.space_after = Pt(12)
    flow.paragraph_format.keep_with_next = True
    fr = flow.add_run(
        "DISCOVER  →  VERIFY  →  QUALIFY  →  ANALYSE  →  SOURCE  →  PRICE  "
        "→  APPROVE  →  EXECUTE  →  LEARN"
    )
    set_run_font(fr, size=9.5, color=PARTNER_NAVY, bold=True)

    add_numbered_step(
        doc,
        numbered_id,
        "Find and verify opportunities. ",
        "The Radar layer monitors tender and export sources and identifies "
        "retenders, corrigenda, repeat-buyer patterns and potentially "
        "under-served opportunities. A public listing is treated as a lead, not "
        "as proof that the opportunity is ready to pursue.",
    )
    add_numbered_step(
        doc,
        numbered_id,
        "Qualify quickly. ",
        "A Fast Kill stage checks the obvious constraints—eligibility, deadline, "
        "EMD, prior experience, buyer risk and product restrictions. Clear "
        "non-starters are removed early; incomplete or ambiguous cases move to a "
        "watchlist instead of being rejected on assumptions.",
    )
    # Named pagination override: section breaks keep the opening page concise
    # while preserving Word-native page geometry and linked running furniture.
    add_partner_section(doc)
    add_numbered_step(
        doc,
        numbered_id,
        "Understand the documents. ",
        "The Deep Read stage extracts specifications, quantities, clauses, "
        "payment terms, penalties, delivery requirements and missing information "
        "from tender PDFs, BOQs, spreadsheets, corrigenda and RFQ documents.",
    )
    add_numbered_step(
        doc,
        numbered_id,
        "Prove that supply is viable. ",
        "The Supplier Engine uses a 5-3-2 discipline: five candidate suppliers, "
        "three different source types and two pieces of quotation proof before "
        "final pricing is considered ready.",
    )
    add_numbered_step(
        doc,
        numbered_id,
        "Build the commercial and compliance view. ",
        "Pricing models the complete landed-cost waterfall rather than using a "
        "single supplier number. Export cases can include EXW, FOB and CIF "
        "scenarios, while compliance remains a clearly labelled draft until an "
        "owner or qualified expert confirms it.",
    )
    add_numbered_step(
        doc,
        numbered_id,
        "Prepare the decision pack. ",
        "The system assembles bid packs, pricing sheets, compliance matrices, "
        "supplier comparisons, risk notes, missing-item lists and owner briefs so "
        "that the decision is made from one coherent evidence bundle.",
    )
    add_numbered_step(
        doc,
        numbered_id,
        "Stop for human approval. ",
        "Before any external commitment, the system creates an approval card that "
        "shows the proposed action, value, evidence, expected benefit, risks, "
        "missing information and recovery path.",
    )
    add_numbered_step(
        doc,
        numbered_id,
        "Track execution and outcome. ",
        "After approval, the Execution Tracker records submissions, responses, "
        "follow-ups, deadlines, receipts and final outcomes. The purpose is to "
        "learn from real results rather than from AI confidence alone.",
    )
    add_partner_section(doc)

    add_heading(doc, "How the system is structured", 1)
    add_body(
        doc,
        "The architecture separates the work into clear layers so that research, "
        "execution, control and accountability do not become mixed together.",
        after=8,
    )

    architecture_rows = [
        (
            "Hermes control plane",
            "Runs the operating rhythm, routes work, manages the Kanban board, "
            "surfaces blockers and approval requests, and produces the owner brief.",
        ),
        (
            "Specialist agents",
            "Separate roles handle discovery, qualification, document analysis, "
            "supplier proof, pricing, compliance, pack building, approvals and "
            "execution tracking.",
        ),
        (
            "Python and browser capture",
            "Perform repeatable evidence capture from known sources, document "
            "parsing, deduplication, validation and scheduled internal jobs.",
        ),
        (
            "Codex artifact runtime",
            "Produces and tests the practical outputs: spreadsheets, PDFs, Word "
            "documents, dashboards, bid packs and export quotation packs.",
        ),
        (
            "ChatGPT research boardroom",
            "Handles broad market research, category discovery and strategic "
            "analysis. Its output is advisory until evidence is verified locally.",
        ),
        (
            "Event and knowledge layer",
            "An append-only event ledger is the canonical record. Registers, "
            "Kanban cards, reports and Drive folders are controlled views of that "
            "record rather than competing sources of truth.",
        ),
        (
            "Policy and approval layer",
            "Policy-as-code, restricted tool interfaces and approval receipts keep "
            "external, financial, legal and compliance actions fail-closed.",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2376, 6984])
    set_table_borders(table)
    header_cells = table.rows[0].cells
    header_cells[0].text = "LAYER"
    header_cells[1].text = "ROLE IN THE OPERATING SYSTEM"
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for cell in header_cells:
        set_cell_fill(cell, LIGHT_GRAY)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=PARTNER_NAVY, bold=True)

    for label, detail in architecture_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = detail
        prevent_row_split(table.rows[-1])
        for idx, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9.8,
                        color=PARTNER_NAVY if idx == 0 else BODY,
                        bold=(idx == 0),
                    )
    # Adding rows changes cell XML; reassert exact geometry.
    set_table_geometry(table, [2376, 6984])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)

    add_heading(doc, "What is genuinely different about it", 1)
    differentiators = [
        (
            "It separates AI judgment from operational proof. ",
            "ChatGPT can discover and recommend, but deterministic tools must "
            "capture evidence before a case progresses.",
        ),
        (
            "It is a state machine, not a chain of prompts. ",
            "Every case moves through explicit statuses, gates and stop conditions.",
        ),
        (
            "Evidence has a defined maturity level. ",
            "A public listing, a downloaded document, a supplier quote and an "
            "approved action are not treated as equivalent evidence.",
        ),
        (
            "Governance is part of the architecture. ",
            "The approval boundary is enforced by policy and receipts rather than "
            "left as a reminder inside an AI prompt.",
        ),
        (
            "The system is auditable and recoverable. ",
            "The event stream can reconstruct working views and shows who or what "
            "changed a case, when it changed and what evidence supported the change.",
        ),
    ]
    for prefix, remainder in differentiators:
        add_bullet(doc, bullet_id, prefix + remainder, bold_prefix=prefix)

    add_heading(doc, "What is automated—and what is deliberately not", 1)
    add_labeled_paragraph(
        doc,
        "Automated or automation-assisted: ",
        "source monitoring, evidence capture, document extraction, deduplication, "
        "scoring, supplier research, draft pricing, draft compliance, artifact "
        "generation, readiness checks, reporting and follow-up monitoring.",
        after=8,
    )
    add_labeled_paragraph(
        doc,
        "Human-controlled: ",
        "supplier or buyer communication, bid and quotation submission, portal "
        "uploads, payments, DSC use, final price, delivery commitments, payment "
        "terms, final classification, origin claims and legal certifications.",
        after=10,
    )

    add_heading(doc, "Current maturity", 1)
    add_body(
        doc,
        "The internal operating architecture is implemented as a supervised local "
        "runtime. The repository contains the workflow definitions, event and data "
        "schemas, source adapters, parsers, policy controls, approval lifecycle, "
        "artifact generators, projection rebuilds and extensive regression and "
        "edge-case tests.",
    )
    add_body(
        doc,
        "I would describe it as a working supervised automation platform—not as a "
        "finished zero-touch tender-submission bot. Credentialed portal operations "
        "and binding commercial or legal actions are intentionally outside the "
        "autonomous tool surface. That is a design choice, not an unfinished safety "
        "check.",
        after=12,
    )

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(4)
    closing.paragraph_format.space_after = Pt(0)
    closing.paragraph_format.line_spacing = 1.15
    set_paragraph_shading_and_border(
        closing, fill=LIGHT_BLUE, border_color=BLUE
    )
    close1 = closing.add_run("In simple terms, ")
    set_run_font(close1, size=11, color=PARTNER_NAVY, bold=True)
    close2 = closing.add_run(
        "I have built the digital operating backbone for a tender and export desk. "
        "It reduces dependence on scattered documents, spreadsheets and individual "
        "memory, while preserving the accountability required for high-stakes "
        "commercial work."
    )
    set_run_font(close2, size=11, color=BODY)

    # Core document metadata.
    props = doc.core_properties
    props.title = "Tender Export OS — Partner Brief"
    props.subject = "Governed automation architecture and workflow overview"
    props.author = "Raghav"
    props.keywords = (
        "Tender Export OS, automation, workflow, human-in-the-loop, "
        "event sourcing, policy-as-code"
    )
    props.comments = "Prepared for a partner-level automation discussion."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
