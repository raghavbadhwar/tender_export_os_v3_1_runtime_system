"""Document parsing for tender/RFQ evidence files."""

from __future__ import annotations

import csv
import json
import re
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from .evidence_store import EvidenceStore, relative, safe_name
from .html_parser import html_to_text
from .table_extractor import extract_html_tables
from .table_classifier import classify_table

DATE_PATTERN = re.compile(
    r"\b(?:\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{2,4})\b",
    re.I,
)
MONEY_PATTERN = re.compile(r"(?:Rs\.?|INR|₹)\s?[\d,]+(?:\.\d+)?|[\d,]+(?:\.\d+)?\s?(?:lakh|lakhs|crore|crores)", re.I)


@dataclass
class ParseResult:
    source_path: str
    parse_status: str
    confidence: str
    text: str = ""
    tables: list[Any] = field(default_factory=list)
    dates: list[str] = field(default_factory=list)
    money_amounts: list[str] = field(default_factory=list)
    extracted_text_path: str = ""
    extracted_json_path: str = ""
    page_count: int | None = None
    notes: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _write_outputs(path: Path, text: str, tables: list[Any], confidence: str, evidence: EvidenceStore | None) -> ParseResult:
    dates = sorted(set(DATE_PATTERN.findall(text)))
    money = sorted(set(match.group(0) for match in MONEY_PATTERN.finditer(text)))
    result = ParseResult(
        source_path=relative(path),
        parse_status="PARSED" if text or tables else "FAILED",
        confidence=confidence if text or tables else "FAILED",
        text=text,
        tables=tables,
        dates=dates,
        money_amounts=money,
    )
    if evidence:
        text_path = evidence.path_for("parsed_text", f"{path.stem}.txt")
        text_path.write_text(text, encoding="utf-8")
        result.extracted_text_path = evidence.record_parsed_text(text_path, relative(path), result.confidence)
        result.extracted_json_path = evidence.write_extracted_json(f"{path.stem}_parsed.json", result.to_dict())
    return result


def parse_html(path: Path, evidence: EvidenceStore | None = None) -> ParseResult:
    html = path.read_text(encoding="utf-8", errors="ignore")
    text = html_to_text(html)
    tables = extract_html_tables(html)
    return _write_outputs(path, text, tables, "HIGH" if text else "FAILED", evidence)


def parse_csv(path: Path, evidence: EvidenceStore | None = None) -> ParseResult:
    rows: list[list[str]] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        for row in csv.reader(f):
            rows.append(row)
    text = "\n".join(" | ".join(cell for cell in row if cell) for row in rows)
    return _write_outputs(path, text, [rows] if rows else [], "HIGH" if rows else "FAILED", evidence)


def parse_excel(path: Path, evidence: EvidenceStore | None = None) -> ParseResult:
    try:
        import pandas as pd

        workbook = pd.read_excel(path, sheet_name=None, dtype=str)
        tables = []
        text_parts = []
        for sheet, frame in workbook.items():
            frame = frame.fillna("")
            rows = [frame.columns.astype(str).tolist()] + frame.astype(str).values.tolist()
            tables.append({"sheet": sheet, "rows": rows})
            text_parts.append(f"Sheet: {sheet}\n" + frame.to_csv(index=False))
        return _write_outputs(path, "\n\n".join(text_parts), tables, "HIGH", evidence)
    except Exception as exc:
        return ParseResult(source_path=relative(path), parse_status="FAILED", confidence="FAILED", notes=str(exc))


def parse_docx(path: Path, evidence: EvidenceStore | None = None) -> ParseResult:
    try:
        from docx import Document

        doc = Document(path)
        text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        tables = []
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                tables.append(rows)
        return _write_outputs(path, "\n".join(text_parts), tables, "HIGH" if text_parts or tables else "FAILED", evidence)
    except Exception as exc:
        return ParseResult(source_path=relative(path), parse_status="FAILED", confidence="FAILED", notes=str(exc))


def parse_pdf(path: Path, evidence: EvidenceStore | None = None) -> ParseResult:
    text_parts: list[str] = []
    tables: list[Any] = []
    page_count: int | None = None
    notes = ""
    try:
        import fitz  # PyMuPDF

        with fitz.open(path) as doc:
            page_count = doc.page_count
            for page in doc:
                text_parts.append(page.get_text("text"))
    except Exception as exc:
        notes = f"PyMuPDF failed: {exc}"
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text:
                    text_parts.append(text)
                for table_index, table in enumerate(page.extract_tables() or [], start=1):
                    rows = table or []
                    tables.append(
                        {
                            "page": page_number,
                            "table_index": table_index,
                            "table_type": classify_table(rows),
                            "rows": rows,
                        }
                    )
    except Exception as exc:
        if not text_parts:
            return ParseResult(source_path=relative(path), parse_status="FAILED", confidence="FAILED", notes=f"{notes}; pdfplumber failed: {exc}")
        notes = f"{notes}; pdfplumber failed: {exc}"

    confidence = "HIGH" if any(text_parts) and tables else "MEDIUM" if any(text_parts) else "LOW"
    result = _write_outputs(path, "\n".join(text_parts), tables, confidence, evidence)
    if not text_parts and not tables:
        result.parse_status = "SCANNED_OR_UNREADABLE_PDF"
        result.confidence = "FAILED"
        result.notes = notes or "No extractable text or tables found; manual review required"
    else:
        result.notes = notes
        result.page_count = page_count
    return result


def parse_zip(path: Path, evidence: EvidenceStore | None = None) -> ParseResult:
    try:
        extract_root = path.parent / f"{safe_name(path.stem)}_unzipped"
        extract_root.mkdir(exist_ok=True)
        child_results: list[dict[str, Any]] = []
        with zipfile.ZipFile(path) as archive:
            for member in archive.namelist():
                if member.endswith("/"):
                    continue
                target = extract_root / safe_name(Path(member).name)
                with archive.open(member) as src, target.open("wb") as dst:
                    dst.write(src.read())
                if target.suffix.lower() in SUPPORTED_SUFFIXES:
                    child_results.append(parse_document(target, evidence).to_dict())
        text = json.dumps(child_results, indent=2, ensure_ascii=False)
        return _write_outputs(path, text, child_results, "MEDIUM" if child_results else "FAILED", evidence)
    except Exception as exc:
        return ParseResult(source_path=relative(path), parse_status="FAILED", confidence="FAILED", notes=str(exc))


SUPPORTED_SUFFIXES = {".pdf", ".xlsx", ".xls", ".csv", ".docx", ".html", ".htm", ".zip"}


def parse_document(path: Path | str, evidence: EvidenceStore | None = None) -> ParseResult:
    document_path = Path(path)
    suffix = document_path.suffix.lower()
    if suffix in {".html", ".htm"}:
        return parse_html(document_path, evidence)
    if suffix == ".csv":
        return parse_csv(document_path, evidence)
    if suffix in {".xlsx", ".xls"}:
        return parse_excel(document_path, evidence)
    if suffix == ".docx":
        return parse_docx(document_path, evidence)
    if suffix == ".pdf":
        return parse_pdf(document_path, evidence)
    if suffix == ".zip":
        return parse_zip(document_path, evidence)
    return ParseResult(source_path=relative(document_path), parse_status="FAILED", confidence="FAILED", notes=f"Unsupported file type: {suffix}")
